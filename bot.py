from __future__ import annotations

import os
import logging
import sys
import tempfile
from pathlib import Path

import httpx
import pytesseract
from docx import Document as DocxDocument
from faster_whisper import WhisperModel
from pdf2image import convert_from_path
from pypdf import PdfReader
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, ContextTypes, filters


TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://127.0.0.1:11434/api/chat")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen3:8b")
MAX_HISTORY_MESSAGES = int(os.getenv("MAX_HISTORY_MESSAGES", "12"))
ALLOWED_TELEGRAM_USER_IDS_RAW = os.getenv("ALLOWED_TELEGRAM_USER_IDS", "")
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()

WHISPER_MODEL_SIZE = os.getenv("WHISPER_MODEL_SIZE", "small")
WHISPER_DEVICE = os.getenv("WHISPER_DEVICE", "cpu")
WHISPER_COMPUTE_TYPE = os.getenv("WHISPER_COMPUTE_TYPE", "int8")

MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", "20"))
MAX_DOCUMENT_CHARS = int(os.getenv("MAX_DOCUMENT_CHARS", "18000"))

OCR_DPI = int(os.getenv("OCR_DPI", "200"))
OCR_LANG = os.getenv("OCR_LANG", "rus+eng")
MAX_OCR_PAGES = int(os.getenv("MAX_OCR_PAGES", "20"))

POPPLER_PATH = os.getenv("POPPLER_PATH", "/opt/homebrew/bin")
TESSERACT_CMD = os.getenv("TESSERACT_CMD", "/opt/homebrew/bin/tesseract")

pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
    stream=sys.stdout,
)
LOGGER = logging.getLogger("telegram_ollama_bot")

USER_HISTORY = {}
STT_MODEL = None


BASE_SYSTEM_PROMPT = """
Ты персональный локальный AI-ассистент Алексея Мирошина.

Контекст пользователя:
- Алексей занимается enterprise IT, инфраструктурой, СХД, AI-решениями, проектами SURF Consulting.
- Частые задачи: деловые письма, follow-up после встреч, тексты для заказчиков, чиновников, презентаций, ТЗ, коммерческих предложений.
- Пиши в первую очередь по-русски, если пользователь явно не просит английский.
- Стиль: деловой, ясный, без воды.
- Если задача техническая — отвечай пошагово и давай команды для macOS/Linux.
- Если задача для высокого чиновника — пиши коротко, статусно, понятно, без перегруза терминами.
- Не выдумывай факты. Если данных мало — честно обозначь допущения.
- Не используй markdown-таблицы без необходимости.
"""


MODES = {
    "default": BASE_SYSTEM_PROMPT,

    "email": BASE_SYSTEM_PROMPT + """
Режим: деловое письмо.
Задача: подготовить аккуратное деловое письмо.
Структура:
1. Тема письма
2. Текст письма
3. Более короткая версия, если уместно
Тон: вежливый, уверенный, без канцелярита.
""",

    "rewrite": BASE_SYSTEM_PROMPT + """
Режим: редактор текста.
Задача: переписать текст лучше.
Сохраняй смысл, убирай воду, усиливай формулировки.
Дай только готовую версию, без длинных объяснений.
""",

    "shorten": BASE_SYSTEM_PROMPT + """
Режим: сжатие текста.
Задача: сделать короткую, сильную и понятную версию.
Структура:
1. Коротко для мессенджера
2. Чуть подробнее для email
3. Один главный тезис
""",

    "vip": BASE_SYSTEM_PROMPT + """
Режим: текст для очень высокого руководителя или чиновника.
Стиль: строго, дорого, спокойно, статусно.
Пиши простым языком, без технической перегрузки.
Фокус: государственная значимость, управляемость, эффект, снижение рисков, масштабирование.
""",

    "surf": BASE_SYSTEM_PROMPT + """
Режим: стиль SURF Consulting.
Стиль: enterprise, зрелый, уверенный, деловой.
Фокус: комплексное решение, практическая польза, внедряемость, ответственность, результат для заказчика.
""",

    "shell": BASE_SYSTEM_PROMPT + """
Режим: помощник по терминалу macOS/Linux.
Задача: давать безопасные и понятные команды.
Всегда:
- объясняй, что делает команда;
- предупреждай перед опасными командами;
- давай проверочные команды после изменений.
""",

    "followup": BASE_SYSTEM_PROMPT + """
Режим: follow-up после встречи.
Структура:
1. Тема письма
2. Короткое письмо
3. Что зафиксировано
4. Следующие шаги
5. Что запросить у заказчика
Тон: деловой, конструктивный, без давления.
""",

    "voice": BASE_SYSTEM_PROMPT + """
Режим: обработка голосовой заметки.
Пользователь прислал расшифровку голосового сообщения.

Задача:
1. Коротко пересказать смысл
2. Выделить задачи
3. Выделить важные детали
4. Если это похоже на встречу — предложить follow-up
5. Если это техническая задача — дать конкретные следующие шаги

Пиши структурированно, коротко и без воды.
""",

    "document": BASE_SYSTEM_PROMPT + """
Режим: анализ документа.
Пользователь прислал файл, из которого извлечен текст.

Задача:
1. Дать краткое резюме документа
2. Выделить ключевые тезисы
3. Выделить важные требования, если они есть
4. Выделить задачи и следующие шаги
5. Выделить риски, вопросы и недостающие данные
6. Если документ похож на ТЗ — отдельно дать блок "Что нужно уточнить у заказчика"
7. Если документ похож на коммерческое предложение — отдельно дать блок "Что важно для продажи"

Пиши структурированно, деловым языком, без воды.
""",
}


def parse_allowed_user_ids(raw_value: str) -> set[int]:
    user_ids = set()

    for token in raw_value.replace(",", " ").split():
        try:
            user_id = int(token)
        except ValueError as exc:
            raise ValueError(
                f"Invalid Telegram user ID in ALLOWED_TELEGRAM_USER_IDS: {token}"
            ) from exc

        if user_id <= 0:
            raise ValueError(
                f"Telegram user IDs must be positive integers: {token}"
            )

        user_ids.add(user_id)

    return user_ids


ALLOWED_TELEGRAM_USER_IDS = parse_allowed_user_ids(ALLOWED_TELEGRAM_USER_IDS_RAW)


def is_user_allowed(user_id: int | None) -> bool:
    if not ALLOWED_TELEGRAM_USER_IDS:
        return True

    return user_id in ALLOWED_TELEGRAM_USER_IDS


def get_update_user_id(update: Update) -> int | None:
    if not update.effective_user:
        return None

    return update.effective_user.id


async def reject_unauthorized(update: Update) -> bool:
    user_id = get_update_user_id(update)

    if is_user_allowed(user_id):
        return False

    username = getattr(update.effective_user, "username", None)
    LOGGER.warning(
        "Unauthorized Telegram access rejected: user_id=%s username=%s",
        user_id,
        username,
    )

    if update.message:
        await update.message.reply_text("Доступ к этому боту ограничен.")

    return True


async def ask_ollama(user_id: int, text: str, mode: str = "default") -> str:
    history = USER_HISTORY.setdefault(user_id, [])
    system_prompt = MODES.get(mode, MODES["default"])

    messages = [{"role": "system", "content": system_prompt}]
    messages.extend(history[-MAX_HISTORY_MESSAGES:])
    messages.append({"role": "user", "content": text})

    payload = {
        "model": OLLAMA_MODEL,
        "messages": messages,
        "stream": False,
        "options": {
            "temperature": 0.4,
            "top_p": 0.9
        }
    }

    async with httpx.AsyncClient(timeout=180) as client:
        response = await client.post(OLLAMA_URL, json=payload)
        response.raise_for_status()
        data = response.json()

    answer = data.get("message", {}).get("content", "").strip()

    if not answer:
        answer = "Модель вернула пустой ответ."

    history.append({"role": "user", "content": text})
    history.append({"role": "assistant", "content": answer})
    USER_HISTORY[user_id] = history[-MAX_HISTORY_MESSAGES:]

    return answer


def get_command_text(update: Update, command: str) -> str:
    text = update.message.text or ""
    return text.replace(f"/{command}", "", 1).strip()


def get_stt_model():
    global STT_MODEL

    if STT_MODEL is None:
        LOGGER.info(
            "Loading Whisper model: model=%s device=%s compute_type=%s",
            WHISPER_MODEL_SIZE,
            WHISPER_DEVICE,
            WHISPER_COMPUTE_TYPE,
        )

        STT_MODEL = WhisperModel(
            WHISPER_MODEL_SIZE,
            device=WHISPER_DEVICE,
            compute_type=WHISPER_COMPUTE_TYPE
        )

        LOGGER.info("Whisper model loaded")

    return STT_MODEL


def transcribe_audio_file(audio_path: str) -> str:
    model = get_stt_model()

    segments, info = model.transcribe(
        audio_path,
        language="ru",
        vad_filter=True,
        beam_size=5
    )

    text_parts = []

    for segment in segments:
        chunk = segment.text.strip()
        if chunk:
            text_parts.append(chunk)

    return " ".join(text_parts).strip()


def extract_text_from_txt(file_path: str) -> str:
    encodings = ["utf-8", "utf-8-sig", "cp1251", "latin-1"]

    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue

    raise ValueError("Не удалось прочитать текстовый файл: неизвестная кодировка")


def extract_text_from_pdf_direct(file_path: str) -> str:
    reader = PdfReader(file_path)
    pages_text = []

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()

        if text:
            pages_text.append(f"\n--- Страница {page_num} ---\n{text}")

    return "\n".join(pages_text).strip()


def extract_text_from_pdf_ocr(file_path: str) -> str:
    LOGGER.info(
        "OCR started: dpi=%s lang=%s max_pages=%s poppler_path=%s tesseract_cmd=%s",
        OCR_DPI,
        OCR_LANG,
        MAX_OCR_PAGES,
        POPPLER_PATH,
        TESSERACT_CMD,
    )

    images = convert_from_path(
        file_path,
        dpi=OCR_DPI,
        fmt="png",
        first_page=1,
        last_page=MAX_OCR_PAGES,
        poppler_path=POPPLER_PATH
    )

    ocr_text_parts = []

    for page_num, image in enumerate(images, start=1):
        LOGGER.info("OCR page %s/%s", page_num, len(images))

        text = pytesseract.image_to_string(
            image,
            lang=OCR_LANG
        ).strip()

        if text:
            ocr_text_parts.append(f"\n--- OCR страница {page_num} ---\n{text}")

    result = "\n".join(ocr_text_parts).strip()

    LOGGER.info("OCR finished")

    return result


def extract_text_from_pdf(file_path: str) -> tuple[str, bool]:
    direct_text = extract_text_from_pdf_direct(file_path)

    if direct_text:
        return direct_text, False

    ocr_text = extract_text_from_pdf_ocr(file_path)

    return ocr_text, True


def extract_text_from_docx(file_path: str) -> str:
    document = DocxDocument(file_path)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                if cell_text:
                    cells.append(cell_text)

            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_text_from_file(file_path: str, filename: str) -> tuple[str, bool]:
    lower_name = filename.lower()

    if lower_name.endswith(".txt") or lower_name.endswith(".md"):
        return extract_text_from_txt(file_path), False

    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_path)

    if lower_name.endswith(".docx"):
        return extract_text_from_docx(file_path), False

    raise ValueError("Поддерживаются только файлы .txt, .md, .pdf и .docx")


def trim_document_text(text: str) -> tuple[str, bool]:
    text = text.strip()

    if len(text) <= MAX_DOCUMENT_CHARS:
        return text, False

    return text[:MAX_DOCUMENT_CHARS], True


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await reject_unauthorized(update):
        return

    await update.message.reply_text(
        "Я локальный AI-бот на Ollama.\n\n"
        "Команды:\n"
        "/status — статус\n"
        "/model — текущая модель\n"
        "/reset — очистить историю\n"
        "/email — деловое письмо\n"
        "/rewrite — переписать текст\n"
        "/shorten — сократить текст\n"
        "/vip — текст для высокого руководителя\n"
        "/surf — стиль SURF Consulting\n"
        "/shell — помощь по терминалу\n"
        "/followup — письмо после встречи\n\n"
        "Голосовые сообщения тоже можно отправлять — я распознаю их локально.\n"
        "Файлы .txt, .md, .pdf и .docx тоже можно отправлять — я извлеку текст и проанализирую.\n"
        "Если PDF является сканом, попробую распознать его через OCR.\n\n"
        "Можно писать и обычным сообщением."
    )


async def myid(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(f"Ваш Telegram ID: {update.effective_user.id}")


async def reset(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await reject_unauthorized(update):
        return

    USER_HISTORY.pop(update.effective_user.id, None)
    await update.message.reply_text("История диалога очищена.")


async def model(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await reject_unauthorized(update):
        return

    await update.message.reply_text(f"Текущая модель: {OLLAMA_MODEL}")


async def status(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await reject_unauthorized(update):
        return

    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.post(
                OLLAMA_URL,
                json={
                    "model": OLLAMA_MODEL,
                    "messages": [
                        {"role": "user", "content": "Ответь одним словом: OK"}
                    ],
                    "stream": False
                }
            )
            response.raise_for_status()

        await update.message.reply_text(
            "Статус: работает\n"
            f"Модель: {OLLAMA_MODEL}\n"
            f"Ollama URL: {OLLAMA_URL}\n"
            f"Whisper: {WHISPER_MODEL_SIZE}, {WHISPER_DEVICE}, {WHISPER_COMPUTE_TYPE}\n"
            f"Файлы: лимит {MAX_FILE_SIZE_MB} MB, анализ до {MAX_DOCUMENT_CHARS} символов\n"
            f"OCR: lang={OCR_LANG}, dpi={OCR_DPI}, max_pages={MAX_OCR_PAGES}\n"
            f"Poppler path: {POPPLER_PATH}\n"
            f"Tesseract cmd: {TESSERACT_CMD}"
        )
    except Exception:
        LOGGER.exception("Status check failed")
        await update.message.reply_text(
            "Статус: ошибка\nНе удалось обратиться к Ollama. Подробности записаны в лог."
        )


async def handle_mode(update: Update, context: ContextTypes.DEFAULT_TYPE, mode: str):
    if await reject_unauthorized(update):
        return

    text = get_command_text(update, mode)

    if not text:
        examples = {
            "email": "Например:\n/email Напиши письмо партнеру: попроси прислать спецификацию Dell до пятницы",
            "rewrite": "Например:\n/rewrite Сделай текст более деловым: ...",
            "shorten": "Например:\n/shorten Сократи этот текст для Telegram: ...",
            "vip": "Например:\n/vip Опиши проект цифрового медицинского профиля для министра",
            "surf": "Например:\n/surf Упакуй этот текст в стиле SURF Consulting: ...",
            "shell": "Например:\n/shell Как проверить, что Telegram-бот запущен и нет дублей?",
            "followup": "Например:\n/followup После встречи обсудили СХД, заказчик ждет сравнение Dell/HPE/NetApp"
        }
        await update.message.reply_text(examples.get(mode, "Пришлите текст после команды."))
        return

    await update.message.chat.send_action(action="typing")

    try:
        answer = await ask_ollama(update.effective_user.id, text, mode=mode)
        await update.message.reply_text(answer)
    except Exception:
        LOGGER.exception(
            "Ollama request failed: user_id=%s mode=%s",
            get_update_user_id(update),
            mode,
        )
        await update.message.reply_text(
            "Ошибка при обращении к Ollama. Подробности записаны в лог."
        )


async def email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await handle_mode(update, context, "email")


async def rewrite(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await handle_mode(update, context, "rewrite")


async def shorten(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await handle_mode(update, context, "shorten")


async def vip(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await handle_mode(update, context, "vip")


async def surf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await handle_mode(update, context, "surf")


async def shell(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await handle_mode(update, context, "shell")


async def followup(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await handle_mode(update, context, "followup")


async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await reject_unauthorized(update):
        return

    await update.message.reply_text("Голосовое получил. Распознаю локально...")

    try:
        voice = update.message.voice

        if not voice:
            await update.message.reply_text("Не вижу голосового сообщения.")
            return

        tg_file = await context.bot.get_file(voice.file_id)

        with tempfile.TemporaryDirectory() as tmpdir:
            audio_path = str(Path(tmpdir) / "voice.ogg")
            await tg_file.download_to_drive(audio_path)

            transcript = transcribe_audio_file(audio_path)

        if not transcript:
            await update.message.reply_text("Не удалось распознать речь.")
            return

        prompt = (
            "Это расшифровка голосового сообщения пользователя.\n\n"
            f"Расшифровка:\n{transcript}\n\n"
            "Обработай эту голосовую заметку: кратко перескажи смысл, "
            "выдели задачи, важные детали и следующие шаги."
        )

        answer = await ask_ollama(update.effective_user.id, prompt, mode="voice")

        final_answer = (
            "Расшифровка:\n"
            f"{transcript}\n\n"
            "Обработка:\n"
            f"{answer}"
        )

        await update.message.reply_text(final_answer)

    except Exception:
        LOGGER.exception(
            "Voice processing failed: user_id=%s",
            get_update_user_id(update),
        )
        await update.message.reply_text(
            "Ошибка при обработке голосового. Подробности записаны в лог."
        )


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await reject_unauthorized(update):
        return

    document = update.message.document

    if not document:
        await update.message.reply_text("Не вижу файла.")
        return

    filename = document.file_name or "document"
    file_size = document.file_size or 0
    file_size_mb = file_size / 1024 / 1024

    if file_size_mb > MAX_FILE_SIZE_MB:
        await update.message.reply_text(
            f"Файл слишком большой: {file_size_mb:.1f} MB.\n"
            f"Лимит сейчас: {MAX_FILE_SIZE_MB} MB."
        )
        return

    supported = filename.lower().endswith((".txt", ".md", ".pdf", ".docx"))

    if not supported:
        await update.message.reply_text(
            "Пока поддерживаю только файлы:\n"
            ".txt\n"
            ".md\n"
            ".pdf\n"
            ".docx"
        )
        return

    await update.message.reply_text(
        f"Файл получил: {filename}\n"
        "Извлекаю текст. Если это скан, включу OCR..."
    )

    try:
        tg_file = await context.bot.get_file(document.file_id)

        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = str(Path(tmpdir) / filename)
            await tg_file.download_to_drive(file_path)

            extracted_text, used_ocr = extract_text_from_file(file_path, filename)

        if not extracted_text:
            await update.message.reply_text(
                "Не удалось извлечь текст из файла даже через OCR. "
                "Возможно, качество скана слишком низкое, документ защищён или это файл без читаемого текста."
            )
            return

        prepared_text, was_trimmed = trim_document_text(extracted_text)

        prompt = (
            f"Пользователь прислал файл: {filename}\n\n"
            "Ниже извлеченный текст документа.\n\n"
            f"{prepared_text}\n\n"
            "Проанализируй документ. Дай краткое резюме, ключевые тезисы, "
            "важные требования, задачи, риски, вопросы и следующие шаги."
        )

        answer = await ask_ollama(update.effective_user.id, prompt, mode="document")

        header = f"Анализ файла: {filename}\n\n"

        if used_ocr:
            header += (
                "Текстовый слой не найден, поэтому я использовал OCR-распознавание скана.\n\n"
            )

        if was_trimmed:
            header += (
                "Важно: документ длинный, я проанализировал первую часть текста. "
                "Для полного анализа позже лучше добавить RAG/индексацию документов.\n\n"
            )

        await update.message.reply_text(header + answer)

    except Exception:
        LOGGER.exception(
            "Document processing failed: user_id=%s filename=%s",
            get_update_user_id(update),
            filename,
        )
        await update.message.reply_text(
            "Ошибка при обработке файла. Подробности записаны в лог."
        )


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await reject_unauthorized(update):
        return

    text = update.message.text.strip()

    if not text:
        return

    await update.message.chat.send_action(action="typing")

    try:
        answer = await ask_ollama(update.effective_user.id, text, mode="default")
        await update.message.reply_text(answer)
    except Exception:
        LOGGER.exception(
            "Ollama request failed: user_id=%s mode=default",
            get_update_user_id(update),
        )
        await update.message.reply_text(
            "Ошибка при обращении к Ollama. Подробности записаны в лог."
        )


def main():
    if not TELEGRAM_TOKEN:
        raise RuntimeError("Не задан TELEGRAM_TOKEN")

    LOGGER.info("Bot starting")
    LOGGER.info("OLLAMA_URL: %s", OLLAMA_URL)
    LOGGER.info("OLLAMA_MODEL: %s", OLLAMA_MODEL)
    LOGGER.info("MAX_HISTORY_MESSAGES: %s", MAX_HISTORY_MESSAGES)
    LOGGER.info(
        "WHISPER: model=%s device=%s compute_type=%s",
        WHISPER_MODEL_SIZE,
        WHISPER_DEVICE,
        WHISPER_COMPUTE_TYPE,
    )
    LOGGER.info(
        "DOCUMENTS: max_file_size=%sMB max_document_chars=%s",
        MAX_FILE_SIZE_MB,
        MAX_DOCUMENT_CHARS,
    )
    LOGGER.info(
        "OCR: lang=%s dpi=%s max_pages=%s poppler_path=%s tesseract_cmd=%s",
        OCR_LANG,
        OCR_DPI,
        MAX_OCR_PAGES,
        POPPLER_PATH,
        TESSERACT_CMD,
    )

    if ALLOWED_TELEGRAM_USER_IDS:
        LOGGER.info(
            "Access control enabled: allowed_user_count=%s",
            len(ALLOWED_TELEGRAM_USER_IDS),
        )
    else:
        LOGGER.warning("Access control disabled: ALLOWED_TELEGRAM_USER_IDS is empty")

    app = Application.builder().token(TELEGRAM_TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("myid", myid))
    app.add_handler(CommandHandler("reset", reset))
    app.add_handler(CommandHandler("model", model))
    app.add_handler(CommandHandler("status", status))

    app.add_handler(CommandHandler("email", email))
    app.add_handler(CommandHandler("rewrite", rewrite))
    app.add_handler(CommandHandler("shorten", shorten))
    app.add_handler(CommandHandler("vip", vip))
    app.add_handler(CommandHandler("surf", surf))
    app.add_handler(CommandHandler("shell", shell))
    app.add_handler(CommandHandler("followup", followup))

    app.add_handler(MessageHandler(filters.VOICE, handle_voice))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))

    LOGGER.info("Bot polling started")

    app.run_polling()


if __name__ == "__main__":
    main()
