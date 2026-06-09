from __future__ import annotations

import logging

import pytesseract
from docx import Document as DocxDocument
from pdf2image import convert_from_path
from pypdf import PdfReader

from app.config import SETTINGS


LOGGER = logging.getLogger("telegram_ollama_bot")
pytesseract.pytesseract.tesseract_cmd = SETTINGS.tesseract_cmd


def extract_text_from_txt(file_path: str) -> str:
    encodings = ["utf-8", "utf-8-sig", "cp1251", "latin-1"]

    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue

    raise ValueError("Не удалось прочитать текстовый файл: неизвестная кодировка")


def extract_text_from_pdf_direct(file_path: str) -> str:
    reader = PdfReader(file_path)
    pages_text = []

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()

        if text:
            pages_text.append(f"\n--- Страница {page_num} ---\n{text}")

    return "\n".join(pages_text).strip()


def extract_text_from_pdf_ocr(file_path: str) -> str:
    LOGGER.info(
        "OCR started: dpi=%s lang=%s max_pages=%s poppler_path=%s tesseract_cmd=%s",
        SETTINGS.ocr_dpi,
        SETTINGS.ocr_lang,
        SETTINGS.max_ocr_pages,
        SETTINGS.poppler_path,
        SETTINGS.tesseract_cmd,
    )

    images = convert_from_path(
        file_path,
        dpi=SETTINGS.ocr_dpi,
        fmt="png",
        first_page=1,
        last_page=SETTINGS.max_ocr_pages,
        poppler_path=SETTINGS.poppler_path
    )

    ocr_text_parts = []

    for page_num, image in enumerate(images, start=1):
        LOGGER.info("OCR page %s/%s", page_num, len(images))

        text = pytesseract.image_to_string(
            image,
            lang=SETTINGS.ocr_lang
        ).strip()

        if text:
            ocr_text_parts.append(f"\n--- OCR страница {page_num} ---\n{text}")

    result = "\n".join(ocr_text_parts).strip()

    LOGGER.info("OCR finished")

    return result


def extract_text_from_pdf(file_path: str) -> tuple[str, bool]:
    direct_text = extract_text_from_pdf_direct(file_path)

    if direct_text:
        return direct_text, False

    ocr_text = extract_text_from_pdf_ocr(file_path)

    return ocr_text, True


def extract_text_from_docx(file_path: str) -> str:
    document = DocxDocument(file_path)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                if cell_text:
                    cells.append(cell_text)

            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_text_from_file(file_path: str, filename: str) -> tuple[str, bool]:
    lower_name = filename.lower()

    if lower_name.endswith(".txt") or lower_name.endswith(".md"):
        return extract_text_from_txt(file_path), False

    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_path)

    if lower_name.endswith(".docx"):
        return extract_text_from_docx(file_path), False

    raise ValueError("Поддерживаются только файлы .txt, .md, .pdf и .docx")


def trim_document_text(text: str, max_chars: int | None = None) -> tuple[str, bool]:
    limit = max_chars if max_chars is not None else SETTINGS.max_document_chars
    text = text.strip()

    if len(text) <= limit:
        return text, False

    return text[:limit], True
